"""
05_targeted_analysis.py
=======================
Three targeted policy analyses:
  1. Elderly Food Access Vulnerability
  2. SNAP / WIC Gap Analysis
  3. Dollar Store Dependence — Current + Tipping Point Risk

Reads  : data/nj_zip_features_v5.csv  +  data/nj_zip_scores_1.csv
Outputs: data/  — CSV files
         plots/ — PNG figures
"""

from pathlib import Path
import warnings
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import importlib
sns = importlib.import_module("seaborn")
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
from sklearn.model_selection import cross_val_score
from sklearn.inspection import permutation_importance
from scipy.stats import spearmanr, mannwhitneyu

warnings.filterwarnings("ignore")

# ═════════════════════════════════════════════════════════════════════════════
# PATHS
# ═════════════════════════════════════════════════════════════════════════════
BASE_DIR  = Path(__file__).resolve().parent
DATA      = BASE_DIR / "data"
PLOTS     = BASE_DIR / "plots"
PLOTS.mkdir(exist_ok=True)

FEATURES_FILE = DATA / "nj_zip_features_v5.csv"
SCORES_FILE   = DATA / "nj_zip_scores_1.csv"

plt.rcParams.update({
    "font.family":       "DejaVu Sans",
    "font.size":         11,
    "axes.spines.top":   False,
    "axes.spines.right": False,
    "figure.dpi":        150,
    "savefig.dpi":       300,
    "savefig.bbox":      "tight",
})

PALETTE = {
    "elderly":   "#8E44AD",
    "snap":      "#27AE60",
    "wic":       "#2980B9",
    "dollar":    "#E67E22",
    "risk":      "#C0392B",
    "safe":      "#ECF0F1",
    "neutral":   "#7F8C8D",
}

# ═════════════════════════════════════════════════════════════════════════════
# LOAD
# ═════════════════════════════════════════════════════════════════════════════
def section(title):
    print(f"\n{'═' * 60}")
    print(f"  {title}")
    print(f"{'═' * 60}")

def save(frame, name):
    path = DATA / name
    frame.to_csv(path, index=False)
    print(f"  → Saved: data/{name}  ({len(frame):,} rows)")

def savefig(name):
    plt.savefig(PLOTS / name)
    plt.close()
    print(f"  → Saved: plots/{name}")

section("LOAD")

df = pd.read_csv(FEATURES_FILE, dtype={"zip": str})
df["zip"] = df["zip"].str.zfill(5)

if SCORES_FILE.exists():
    scores = pd.read_csv(SCORES_FILE, dtype={"zip": str})
    scores["zip"] = scores["zip"].str.zfill(5)
    pred_cols = [c for c in scores.columns if c.startswith("predicted_")
                 or c.startswith("typo_prob_")]
    model_cols = [c for c in ["desert_probability", "predicted_desert",
                               "predicted_desert_tuned", "predicted_typology"]
                  if c in scores.columns]
    cols_to_merge = ["zip"] + model_cols + pred_cols
    new_cols = [c for c in cols_to_merge if c not in df.columns or c == "zip"]
    df = df.merge(scores[new_cols], on="zip", how="left")

print(f"  Loaded: {len(df):,} ZIPs × {len(df.columns):,} columns")

# Shared derived columns used across analyses
df["pct_white"] = (
    df["Population Non-Hispanic White"] / df["population"].replace(0, np.nan)
    if "Population Non-Hispanic White" in df.columns else np.nan
)

# ═════════════════════════════════════════════════════════════════════════════
# ANALYSIS 1 — ELDERLY FOOD ACCESS VULNERABILITY
# ═════════════════════════════════════════════════════════════════════════════
section("ANALYSIS 1 — ELDERLY FOOD ACCESS VULNERABILITY")

# ── 1a. Elderly concentration threshold ──────────────────────────────────────
# "Elderly-heavy" = top tertile of pct_elderly
elderly_thresh  = df["pct_elderly"].quantile(0.67)
df["is_elderly_heavy"] = (df["pct_elderly"] >= elderly_thresh).astype(int)
n_elderly = df["is_elderly_heavy"].sum()
print(f"\n  Elderly-heavy ZIPs (top tertile, ≥{elderly_thresh:.1f}%): {n_elderly}")

# ── 1b. Elderly-specific vulnerability score ─────────────────────────────────
# Weighted composite targeting the barriers most relevant to older adults:
#   physical distance  30%
#   no vehicle         25%
#   food swamp ratio   20%   (elderly less mobile, can't avoid nearest options)
#   poverty            15%
#   low SNAP access    10%

def pct_rank(s):
    return s.rank(pct=True)

snap_density = df["snap_stores"] / df["population"].replace(0, np.nan) * 10000

df["elderly_food_vuln_score"] = (
    pct_rank(df["nearest_supermarket_miles"])      * 0.30
    + pct_rank(df["pct_no_vehicle"])               * 0.25
    + pct_rank(df["swamp_score_continuous"])        * 0.20
    + pct_rank(df["pct_poverty"])                  * 0.15
    + (1 - pct_rank(snap_density.fillna(0)))       * 0.10
) * 100

# ── 1c. Elderly risk tier ────────────────────────────────────────────────────
df["elderly_risk_tier"] = pd.qcut(
    df["elderly_food_vuln_score"],
    q=4,
    labels=["Low Risk", "Moderate Risk", "High Risk", "Critical Risk"]
)

# ── 1d. Summary: elderly-heavy ZIPs by risk tier ─────────────────────────────
elderly_df = df[df["is_elderly_heavy"] == 1].copy()

tier_summary = (
    elderly_df.groupby("elderly_risk_tier", observed=True)
    .agg(
        zip_count               = ("zip",                       "count"),
        population              = ("population",                "sum"),
        avg_pct_elderly         = ("pct_elderly",               "mean"),
        avg_pct_no_vehicle      = ("pct_no_vehicle",            "mean"),
        avg_nearest_super_mi    = ("nearest_supermarket_miles", "mean"),
        avg_swamp_score         = ("swamp_score_continuous",    "mean"),
        avg_pct_poverty         = ("pct_poverty",               "mean"),
        avg_median_income       = ("median_income",             "mean"),  # ← this line
        avg_snap_stores         = ("snap_stores",               "mean"),
        avg_wic_stores          = ("wic_stores",                "mean"),
    )
    .reset_index()
)

ELDERLY_HEALTH = [c for c in [
    "Coronary Heart Disease % (Adults)",
    "Diabetes % (Adults)",
    "Any Disability % (Adults)",
    "All Teeth Lost % (Adults 65+)",
    "High Cholesterol % (Adults)",
    "Stroke % (Adults)",
    "Depression % (Adults)",
    "Poor Physical Health % (Adults)",
] if c in df.columns]

for col in ELDERLY_HEALTH:
    tier_summary[f"avg_{col.replace(' % (Adults)', '').replace(' % (Adults 65+)', '_65plus').replace(' ', '_').lower()}"] = \
        elderly_df.groupby("elderly_risk_tier", observed=True)[col].mean().values

print(f"\n  Elderly-heavy ZIPs by risk tier:\n")
print(tier_summary[["elderly_risk_tier", "zip_count", "population",
                     "avg_pct_elderly", "avg_pct_no_vehicle",
                     "avg_nearest_super_mi", "avg_pct_poverty"]].to_string(index=False))
save(tier_summary, "elderly_risk_tier_summary.csv")

# ── 1e. Top 30 critical elderly ZIPs ─────────────────────────────────────────
ELDERLY_COLS = [c for c in [
    "zip", "county", "municipality", "population",
    "pct_elderly", "pct_no_vehicle", "pct_poverty",
    "avg_median_income", "median_income",
    "nearest_supermarket_miles", "supermarket",
    "snap_stores", "wic_stores",
    "swamp_score_continuous", "mrfei",
    "access_typology", "elderly_food_vuln_score", "elderly_risk_tier",
    "Coronary Heart Disease % (Adults)",
    "Any Disability % (Adults)",
    "All Teeth Lost % (Adults 65+)",
    "Stroke % (Adults)",
] if c in df.columns]

critical_elderly = (
    df[df["is_elderly_heavy"] == 1]
    .sort_values("elderly_food_vuln_score", ascending=False)
    .head(30)[ELDERLY_COLS]
)
print(f"\n  Top 30 critical elderly ZIPs:\n")
print(critical_elderly[["zip", "county", "municipality", "pct_elderly",
                          "nearest_supermarket_miles", "pct_no_vehicle",
                          "elderly_food_vuln_score", "elderly_risk_tier"]].to_string(index=False))
save(critical_elderly, "elderly_critical_zips.csv")

# ── 1f. Elderly vs non-elderly food environment comparison ───────────────────
print(f"\n  Elderly-heavy vs non-elderly-heavy food environment comparison:\n")
compare_cols = [
    "nearest_supermarket_miles", "pct_no_vehicle",
    "swamp_score_continuous", "snap_stores", "wic_stores",
    "supermarket", "pct_poverty",
]
compare_rows = []
for col in [c for c in compare_cols if c in df.columns]:
    hi  = df.loc[df["is_elderly_heavy"] == 1, col].dropna()
    lo  = df.loc[df["is_elderly_heavy"] == 0, col].dropna()
    _, p = mannwhitneyu(hi, lo, alternative="two-sided")
    stars = "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else "ns"
    compare_rows.append({
        "metric":           col,
        "elderly_heavy_mean":  round(hi.mean(), 3),
        "non_elderly_mean":    round(lo.mean(), 3),
        "diff":                round(hi.mean() - lo.mean(), 3),
        "significance":        stars,
    })
compare_df = pd.DataFrame(compare_rows)
print(compare_df.to_string(index=False))
save(compare_df, "elderly_food_environment_comparison.csv")

# ── 1g. County-level elderly summary ─────────────────────────────────────────
elderly_county = (
    df[df["is_elderly_heavy"] == 1]
    .groupby("county")
    .agg(
        n_elderly_heavy_zips    = ("zip",                        "count"),
        pct_critical_risk       = ("elderly_risk_tier",
                                   lambda x: (x == "Critical Risk").mean() * 100),
        avg_elderly_vuln_score  = ("elderly_food_vuln_score",    "mean"),
        avg_pct_elderly         = ("pct_elderly",                "mean"),
        avg_nearest_super_mi    = ("nearest_supermarket_miles",  "mean"),
        avg_pct_no_vehicle      = ("pct_no_vehicle",             "mean"),
        total_elderly_pop       = ("population",                 "sum"),
    )
    .reset_index()
    .sort_values("avg_elderly_vuln_score", ascending=False)
)
print(f"\n  Counties ranked by elderly food vulnerability:\n")
print(elderly_county.to_string(index=False))
save(elderly_county, "elderly_county_summary.csv")

# ── 1h. Spearman: pct_elderly × food access metrics ─────────────────────────
print(f"\n  Spearman correlations: % elderly × food access metrics:\n")
spear_rows = []
for col in [c for c in compare_cols + ELDERLY_HEALTH if c in df.columns]:
    r, p = spearmanr(df["pct_elderly"].dropna(),
                     df[col].reindex(df["pct_elderly"].dropna().index))
    spear_rows.append({
        "metric": col,
        "spearman_r": round(r, 3),
        "p_value":    round(p, 4),
        "sig":        "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else "ns",
    })
spear_df = pd.DataFrame(spear_rows).sort_values("spearman_r", ascending=False)
print(spear_df.to_string(index=False))
save(spear_df, "elderly_spearman_correlations.csv")

# ── 1i. Plots ─────────────────────────────────────────────────────────────────
fig, axes = plt.subplots(1, 3, figsize=(16, 5))

# Left: elderly vuln score distribution — elderly-heavy vs not
ax = axes[0]
df[df["is_elderly_heavy"] == 0]["elderly_food_vuln_score"].hist(
    bins=25, ax=ax, alpha=0.6, color=PALETTE["neutral"], label="Non-Elderly Heavy")
df[df["is_elderly_heavy"] == 1]["elderly_food_vuln_score"].hist(
    bins=25, ax=ax, alpha=0.7, color=PALETTE["elderly"], label="Elderly Heavy")
ax.set_xlabel("Elderly Food Vulnerability Score")
ax.set_ylabel("ZIP Count")
ax.set_title("Vulnerability Score Distribution\nElderly vs Non-Elderly Heavy ZIPs",
             fontweight="bold")
ax.legend()

# Middle: risk tier bar — elderly-heavy ZIPs only
ax = axes[1]
tier_counts = elderly_df["elderly_risk_tier"].value_counts().sort_index()
colors_tier = [PALETTE["snap"], PALETTE["wic"], PALETTE["dollar"], PALETTE["risk"]]
tier_counts.plot(kind="bar", ax=ax, color=colors_tier, edgecolor="white", width=0.7)
ax.set_xlabel("")
ax.set_ylabel("ZIP Count")
ax.set_title("Risk Tier Distribution\n(Elderly-Heavy ZIPs Only)", fontweight="bold")
ax.set_xticklabels(tier_counts.index.tolist(), rotation=20, ha="right")

# Right: scatter — % elderly × nearest supermarket, sized by % no vehicle
ax = axes[2]
sc = ax.scatter(
    df["pct_elderly"],
    df["nearest_supermarket_miles"],
    c=df["pct_no_vehicle"],
    cmap="YlOrRd",
    s=30, alpha=0.6, edgecolors="none"
)
plt.colorbar(sc, ax=ax, label="% No Vehicle")
ax.set_xlabel("% Elderly Population")
ax.set_ylabel("Distance to Nearest Supermarket (mi)")
ax.set_title("Elderly Concentration vs.\nSupermarket Distance", fontweight="bold")
ax.axvline(elderly_thresh, color=PALETTE["elderly"], linestyle="--",
           lw=1.2, label=f"Elderly threshold ({elderly_thresh:.0f}%)")
ax.legend(fontsize=9)

plt.suptitle("Elderly Food Access Vulnerability Analysis", fontweight="bold",
             fontsize=13, y=1.02)
plt.tight_layout()
savefig("30_elderly_food_vulnerability.png")

# Heatmap: elderly health outcomes by risk tier
if ELDERLY_HEALTH:
    health_by_tier = elderly_df.groupby(
        "elderly_risk_tier", observed=True)[ELDERLY_HEALTH].mean()
    health_z = (health_by_tier - health_by_tier.mean()) / health_by_tier.std()
    short = [c.replace(" % (Adults)", "").replace(" % (Adults 65+)", " 65+")
             for c in ELDERLY_HEALTH]

    fig, ax = plt.subplots(figsize=(12, 5))
    sns.heatmap(
        health_z.T, annot=health_by_tier.T.round(1), fmt=".1f",
        cmap="RdYlGn_r", center=0, ax=ax,
        xticklabels=health_by_tier.index.tolist(),
        yticklabels=short,
        linewidths=0.4, linecolor="white",
    )
    ax.set_title("Health Outcomes by Elderly Risk Tier\n"
                 "(Annotated: raw %, color: z-score)", fontweight="bold")
    plt.xticks(rotation=15, ha="right")
    plt.tight_layout()
    savefig("31_elderly_health_by_risk_tier.png")


# ═════════════════════════════════════════════════════════════════════════════
# ANALYSIS 2 — SNAP / WIC GAP ANALYSIS
# ═════════════════════════════════════════════════════════════════════════════
section("ANALYSIS 2 — SNAP / WIC GAP ANALYSIS")

# ── 2a. Need and supply measures ─────────────────────────────────────────────
# Need proxy: pct_snap + pct_poverty (normalized to 0-1)
# Supply: outlet density per 10k population

pop = df["population"].replace(0, np.nan)

df["snap_outlets_per_10k"]       = df["snap_stores"]        / pop * 10000
df["snap_super_per_10k"]         = df["snap_supermarkets"]  / pop * 10000
df["wic_outlets_per_10k"]        = df["wic_stores"]         / pop * 10000

# Need score: high pct_snap + high pct_poverty = high need
df["snap_need_score"] = (
    df["pct_snap"].rank(pct=True) * 0.60
    + df["pct_poverty"].rank(pct=True) * 0.40
)

# ── 2b. Gap scores ───────────────────────────────────────────────────────────
# Gap = high need + low supply
# Rank need high, rank supply low — both pointing toward "underserved"

df["snap_gap_score"] = (
    df["snap_need_score"].rank(pct=True)
    + (1 - df["snap_outlets_per_10k"].fillna(0).rank(pct=True))
) / 2 * 100

df["snap_super_gap_score"] = (
    df["snap_need_score"].rank(pct=True)
    + (1 - df["snap_super_per_10k"].fillna(0).rank(pct=True))
) / 2 * 100

df["wic_gap_score"] = (
    df["snap_need_score"].rank(pct=True)          # WIC need proxied by same poverty/snap
    + (1 - df["wic_outlets_per_10k"].fillna(0).rank(pct=True))
) / 2 * 100

# Combined gap — worst of SNAP and WIC gaps
df["snap_wic_combined_gap"] = (
    df["snap_gap_score"] * 0.40
    + df["snap_super_gap_score"] * 0.35
    + df["wic_gap_score"] * 0.25
)

# ── 2c. Gap flags ─────────────────────────────────────────────────────────────
gap_thresh_snap = df["snap_gap_score"].quantile(0.75)
gap_thresh_wic  = df["wic_gap_score"].quantile(0.75)
gap_thresh_comb = df["snap_wic_combined_gap"].quantile(0.75)

df["is_snap_gap_zip"]     = (df["snap_gap_score"]        >= gap_thresh_snap).astype(int)
df["is_wic_gap_zip"]      = (df["wic_gap_score"]         >= gap_thresh_wic).astype(int)
df["is_combined_gap_zip"] = (df["snap_wic_combined_gap"] >= gap_thresh_comb).astype(int)

print(f"\n  Gap thresholds (75th percentile):")
print(f"    SNAP gap         : {gap_thresh_snap:.1f}  →  {df['is_snap_gap_zip'].sum()} ZIPs flagged")
print(f"    WIC gap          : {gap_thresh_wic:.1f}  →  {df['is_wic_gap_zip'].sum()} ZIPs flagged")
print(f"    Combined gap     : {gap_thresh_comb:.1f}  →  {df['is_combined_gap_zip'].sum()} ZIPs flagged")

# ── 2d. Ranked list ───────────────────────────────────────────────────────────
GAP_COLS = [c for c in [
    "zip", "county", "municipality", "median_income", "population",
    "pct_poverty", "pct_snap", "pct_elderly",
    "snap_stores", "snap_supermarkets", "wic_stores",
    "snap_outlets_per_10k", "snap_super_per_10k", "wic_outlets_per_10k",
    "snap_need_score", "snap_gap_score", "snap_super_gap_score",
    "wic_gap_score", "snap_wic_combined_gap",
    "is_snap_gap_zip", "is_wic_gap_zip", "is_combined_gap_zip",
    "access_typology", "nearest_supermarket_miles",
    "is_elderly_heavy", "elderly_food_vuln_score",
] if c in df.columns]

gap_ranked = df[GAP_COLS].sort_values("snap_wic_combined_gap", ascending=False)

print(f"\n  Top 30 underserved SNAP/WIC ZIPs:\n")
print(gap_ranked.head(30)[["zip", "county", "municipality", "median_income",
                             "pct_poverty", "pct_snap",
                             "snap_outlets_per_10k", "wic_outlets_per_10k",
                             "snap_wic_combined_gap",
                             "is_combined_gap_zip"]].to_string(index=False))
save(gap_ranked, "snap_wic_gap_ranked.csv")

# ── 2e. Elderly + SNAP/WIC gap overlap ───────────────────────────────────────
dual_gap = df[(df["is_elderly_heavy"] == 1) & (df["is_combined_gap_zip"] == 1)]
print(f"\n  ZIPs that are BOTH elderly-heavy AND have a SNAP/WIC gap: {len(dual_gap)}")
print(f"  These are highest-priority intervention targets.\n")
dual_gap_out = dual_gap[GAP_COLS].sort_values("snap_wic_combined_gap", ascending=False)
print(dual_gap_out[["zip", "county", "municipality", "pct_elderly",
                      "pct_poverty", "snap_wic_combined_gap"]].to_string(index=False))
save(dual_gap_out, "elderly_snap_wic_dual_gap.csv")

# ── 2f. County gap summary ────────────────────────────────────────────────────
gap_county = (
    df.groupby("county")
    .agg(
        n_zips              = ("zip",                   "count"),
        n_snap_gap          = ("is_snap_gap_zip",       "sum"),
        n_wic_gap           = ("is_wic_gap_zip",        "sum"),
        n_combined_gap      = ("is_combined_gap_zip",   "sum"),
        avg_snap_gap        = ("snap_gap_score",        "mean"),
        avg_wic_gap         = ("wic_gap_score",         "mean"),
        avg_combined_gap    = ("snap_wic_combined_gap", "mean"),
        avg_pct_poverty     = ("pct_poverty",           "mean"),
        avg_pct_snap        = ("pct_snap",              "mean"),
        avg_snap_per_10k    = ("snap_outlets_per_10k",  "mean"),
        avg_wic_per_10k     = ("wic_outlets_per_10k",   "mean"),
        n_elderly_dual_gap  = ("zip", lambda x:
            ((df.loc[x.index, "is_elderly_heavy"] == 1) &
             (df.loc[x.index, "is_combined_gap_zip"] == 1)).sum()),
    )
    .reset_index()
    .sort_values("avg_combined_gap", ascending=False)
)

gap_county["pct_zips_combined_gap"] = (
    gap_county["n_combined_gap"] / gap_county["n_zips"] * 100
).round(1)

print(f"\n  County SNAP/WIC gap ranking:\n")
print(gap_county[["county", "n_zips", "n_combined_gap", "pct_zips_combined_gap",
                   "avg_combined_gap", "avg_pct_poverty",
                   "avg_snap_per_10k", "avg_wic_per_10k",
                   "n_elderly_dual_gap"]].to_string(index=False))
save(gap_county, "snap_wic_gap_by_county.csv")

# ── 2g. Plots ─────────────────────────────────────────────────────────────────
fig, axes = plt.subplots(1, 3, figsize=(17, 5))

# Left: scatter — SNAP need vs SNAP outlets per 10k
ax = axes[0]
colors_gap = df["is_combined_gap_zip"].map({0: PALETTE["snap"], 1: PALETTE["risk"]})
ax.scatter(df["snap_need_score"], df["snap_outlets_per_10k"].clip(upper=50),
           c=colors_gap, alpha=0.55, s=35, edgecolors="none")
ax.set_xlabel("SNAP Need Score (poverty + SNAP use)")
ax.set_ylabel("SNAP Outlets per 10k Population")
ax.set_title("SNAP Need vs. Supply\n(Red = Combined Gap ZIP)", fontweight="bold")
from matplotlib.patches import Patch
ax.legend(handles=[Patch(color=PALETTE["risk"],  label="Gap ZIP"),
                   Patch(color=PALETTE["snap"],  label="Adequate")], loc="upper right")

# Middle: gap score distribution by county (top 10)
ax = axes[1]
top10_counties = gap_county.head(10)
ax.barh(top10_counties["county"], top10_counties["avg_combined_gap"],
        color=PALETTE["risk"], edgecolor="white")
ax.set_xlabel("Mean Combined SNAP/WIC Gap Score")
ax.set_title("Top 10 Counties by\nSNAP/WIC Gap Score", fontweight="bold")
ax.invert_yaxis()

# Right: elderly + gap overlap bar
ax = axes[2]
overlap_data = {
    "Elderly Heavy\nOnly":          (df["is_elderly_heavy"] == 1).sum() - len(dual_gap),
    "SNAP/WIC Gap\nOnly":           (df["is_combined_gap_zip"] == 1).sum() - len(dual_gap),
    "Both (Priority\nTargets)":     len(dual_gap),
    "Neither":                      ((df["is_elderly_heavy"] == 0) &
                                     (df["is_combined_gap_zip"] == 0)).sum(),
}
ax.bar(range(len(overlap_data)), list(overlap_data.values()),
       color=[PALETTE["elderly"], PALETTE["snap"], PALETTE["risk"], PALETTE["neutral"]],
       edgecolor="white")
ax.set_xticks(range(len(overlap_data)))
ax.set_xticklabels(list(overlap_data.keys()), fontsize=9)
ax.set_ylabel("ZIP Count")
ax.set_title("Elderly × SNAP/WIC Gap Overlap", fontweight="bold")

plt.suptitle("SNAP / WIC Gap Analysis", fontweight="bold", fontsize=13, y=1.02)
plt.tight_layout()
savefig("32_snap_wic_gap_analysis.png")

# Heatmap: gap scores by county
pivot_gap = gap_county.set_index("county")[["avg_snap_gap", "avg_wic_gap",
                                             "avg_combined_gap"]].sort_values("avg_combined_gap")
fig, ax = plt.subplots(figsize=(8, 10))
sns.heatmap(pivot_gap, annot=True, fmt=".1f", cmap="YlOrRd",
            linewidths=0.4, linecolor="white", ax=ax,
            xticklabels=["SNAP Gap", "WIC Gap", "Combined Gap"])
ax.set_title("SNAP/WIC Gap Scores by County", fontweight="bold")
plt.tight_layout()
savefig("33_snap_wic_gap_county_heatmap.png")

# ═════════════════════════════════════════════════════════════════════════════
# FINAL SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
section("SUMMARY")

print(f"""
  ELDERLY ANALYSIS
  ─────────────────────────────────────────────────────
  Elderly-heavy ZIPs (≥{elderly_thresh:.0f}% elderly)    : {n_elderly}
  Critical risk elderly ZIPs                    : {(elderly_df['elderly_risk_tier'] == 'Critical Risk').sum()}
  data/elderly_critical_zips.csv
  data/elderly_risk_tier_summary.csv
  data/elderly_county_summary.csv
  data/elderly_snap_wic_dual_gap.csv            ← elderly + SNAP/WIC overlap
  plots/30_elderly_food_vulnerability.png
  plots/31_elderly_health_by_risk_tier.png

  SNAP / WIC GAP ANALYSIS
  ─────────────────────────────────────────────────────
  SNAP gap ZIPs                                 : {df['is_snap_gap_zip'].sum()}
  WIC gap ZIPs                                  : {df['is_wic_gap_zip'].sum()}
  Combined gap ZIPs                             : {df['is_combined_gap_zip'].sum()}
  Elderly + SNAP/WIC dual gap                   : {len(dual_gap)}
  data/snap_wic_gap_ranked.csv
  data/snap_wic_gap_by_county.csv
  plots/32_snap_wic_gap_analysis.png
  plots/33_snap_wic_gap_county_heatmap.png

  DOLLAR STORE DEPENDENCE
  ─────────────────────────────────────────────────────
  data/dollar_store_currently_dominant.csv
  data/dollar_store_tipping_point.csv
  data/dollar_store_high_risk_both.csv
  data/triple_risk_zips.csv                     ← highest priority targets
  data/dollar_store_county_summary.csv
  plots/34_dollar_store_dependence.png
  plots/35_dollar_tipping_point_importance.png
""")

# ═════════════════════════════════════════════════════════════════════════════
# EXPORT TO WORD
# ═════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def df_to_word_table(doc, df, style="Table Grid"):
    """Write a pandas DataFrame into a Word table with a bold header row."""
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = style

    # Header row — bold
    for j, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col_name)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for i, (_, row) in enumerate(df.iterrows()):
        for j, value in enumerate(row):
            table.rows[i + 1].cells[j].text = (
                "" if pd.isna(value) else str(round(value, 3) if isinstance(value, float) else value)
            )

doc = Document()
doc.add_heading("NJ ZIP Food Access — Targeted Analysis", 0)

# ── Analysis 1: Elderly Food Access Vulnerability ───────────────────────────
doc.add_heading("Analysis 1 — Elderly Food Access Vulnerability", level=1)

doc.add_heading("Risk Tier Summary (Elderly-Heavy ZIPs)", level=2)
df_to_word_table(doc, tier_summary[[
    "elderly_risk_tier", "zip_count", "population",
    "avg_pct_elderly", "avg_median_income", "avg_pct_no_vehicle",
    "avg_nearest_super_mi", "avg_pct_poverty"
]].round(2))

doc.add_paragraph()  # spacer
doc.add_heading("Top 30 Critical Elderly ZIPs", level=2)
df_to_word_table(doc, critical_elderly[[
    "zip", "county", "municipality", "pct_elderly", "median_income",  # ← not avg_
    "nearest_supermarket_miles", "pct_no_vehicle",
    "elderly_food_vuln_score", "elderly_risk_tier"
]].round(2))

doc.add_paragraph()
doc.add_heading("Elderly vs Non-Elderly Food Environment Comparison", level=2)
df_to_word_table(doc, compare_df.round(3))

doc.add_paragraph()
doc.add_heading("Counties Ranked by Elderly Food Vulnerability", level=2)
df_to_word_table(doc, elderly_county.round(2))

doc.add_paragraph()
doc.add_heading("Spearman Correlations: % Elderly × Food Access Metrics", level=2)
df_to_word_table(doc, spear_df.round(3))

# ── Analysis 2: SNAP / WIC Gap Analysis ─────────────────────────────────────
doc.add_page_break()
doc.add_heading("Analysis 2 — SNAP / WIC Gap Analysis", level=1)

doc.add_heading("Top 30 Underserved SNAP/WIC ZIPs", level=2)
df_to_word_table(doc, gap_ranked.head(30)[[
    "zip", "county", "municipality",
    "pct_poverty", "median_income", "pct_snap",
    "snap_outlets_per_10k", "wic_outlets_per_10k",
    "snap_wic_combined_gap", "is_combined_gap_zip"
]].round(2))

df_to_word_table(doc, dual_gap_out[[
    "zip", "county", "municipality", "pct_elderly", "median_income",
    "pct_poverty", "snap_wic_combined_gap"
]].round(2))

doc.add_paragraph()
doc.add_heading("County SNAP/WIC Gap Ranking", level=2)
df_to_word_table(doc, gap_county[[
    "county", "n_zips", "n_combined_gap", "pct_zips_combined_gap",
    "avg_combined_gap", "avg_pct_poverty",
    "avg_snap_per_10k", "avg_wic_per_10k", "n_elderly_dual_gap"
]].round(2))

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = BASE_DIR / "data" / "targeted_analysis_report.docx"
doc.save(output_path)
print(f"\n  → Word report saved: {output_path}")