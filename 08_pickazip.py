# zip_profile.py
# Usage: python zip_profile.py --zip 08103
#        python zip_profile.py --zip 08103 --out outputs/08103_profile.docx

import argparse
import json
from pathlib import Path
from datetime import date
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).parent
SCORES_PATH = HERE / "outputs" / "nj_zip_scores.csv"
HEALTH_PATH = HERE / "outputs" / "nj_health_places.csv"
META_PATH   = HERE / "outputs" / "pipeline_metadata.json"

# ── Colour palette ────────────────────────────────────────────────
C_NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_GREY  = RGBColor(0x59, 0x59, 0x59)
C_RED   = RGBColor(0xC0, 0x39, 0x2B)
C_GREEN = RGBColor(0x20, 0x7A, 0x39)
C_AMBER = RGBColor(0xD4, 0x8E, 0x00)

def _shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def _heading(doc, text, size=14, color=C_NAVY, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def _body(doc, text, size=10, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(5)
    return p

def _kv_section(doc, title, rows, header_hex="1A3A5C"):
    """Compact two-column table for a section."""
    _heading(doc, title, size=11)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for i, (k, v, flag) in enumerate(rows):
        fill = "FFF3CD" if flag == "warn" else \
               "FFEBEE" if flag == "alert" else \
               "E8F5E9" if flag == "good" else \
               ("F2F2F2" if i % 2 == 0 else "FFFFFF")
        for j, (cell, text) in enumerate([(table.rows[i].cells[0], k),
                                           (table.rows[i].cells[1], v)]):
            cell.width = Inches(3.0 if j == 0 else 5.5)
            _shade_cell(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.bold = (j == 0)
    doc.add_paragraph()

def _percentile_note(value, series, label):
    """Return a string showing where this ZIP sits in the state distribution."""
    pct = (series < value).mean() * 100
    return f"{label} — state percentile: {pct:.0f}th"

def _flag(value, low_thresh, high_thresh, invert=False):
    """Return a flag string based on thresholds."""
    if invert:
        if value <= low_thresh: return "alert"
        if value <= high_thresh: return "warn"
        return "good"
    else:
        if value >= high_thresh: return "alert"
        if value >= low_thresh:  return "warn"
        return "good"

def build_profile(zip_code: str, out_path: Path):
    # ── Load data ─────────────────────────────────────────────────
    scores = pd.read_csv(SCORES_PATH, dtype={"zip": str}, low_memory=False)
    scores["zip"] = scores["zip"].str.zfill(5)

    row = scores[scores["zip"] == zip_code]
    if row.empty:
        raise ValueError(f"ZIP code {zip_code} not found in {SCORES_PATH}")
    row = row.iloc[0]

    # Load health data if available
    health_row = None
    if HEALTH_PATH.exists():
        health = pd.read_csv(HEALTH_PATH, dtype={"zip": str}, low_memory=False)
        health["zip"] = health["zip"].str.zfill(5)
        h = health[health["zip"] == zip_code]
        if not h.empty:
            health_row = h.iloc[0]

    # After loading scores and health data, add:
    gtfs_path = HERE / "outputs" / f"gtfs_data_{zip_code}.json"
    gtfs = {}
    if gtfs_path.exists():
        gtfs = json.loads(gtfs_path.read_text())

    # Load metadata
    meta = {}
    if META_PATH.exists():
        meta = json.loads(META_PATH.read_text())

    # ── Helper to safely get a value ─────────────────────────────
    def g(col, fmt=None, default="N/A"):
        v = row.get(col, np.nan)
        if pd.isna(v):
            return default
        return fmt.format(v) if fmt else v

    def hg(col, fmt=None, default="N/A"):
        if health_row is None:
            return default
        v = health_row.get(col, np.nan)
        if pd.isna(v):
            return default
        try:
            v = float(v)
        except (ValueError, TypeError):
            return str(v)
        return fmt.format(v) if fmt else str(v)

    # ── Risk classification ───────────────────────────────────────
    prob = float(row.get("desert_probability", 0))
    if prob >= 0.75:
        risk_label, risk_color = "HIGH RISK", C_RED
    elif prob >= 0.40:
        risk_label, risk_color = "MODERATE RISK", C_AMBER
    else:
        risk_label, risk_color = "LOW RISK", C_GREEN

    typology = row.get("food_access_type", "Unknown")

    # ── Build document ────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    today = date.today().strftime("%B %d, %Y")
    acs_yr = meta.get("acs_year", 2022)

    # Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(f"Food Access Profile — ZIP Code {zip_code}")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = C_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"New Jersey Food Desert Risk Assessment  |  Generated: {today}"
                ).font.size = Pt(10)

    # Risk badge
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = badge.add_run(f"  ◆  FOOD ACCESS RISK: {risk_label}  ◆  "
                       f"Desert Probability: {prob:.1%}  ◆  "
                       f"Access Type: {typology}  ◆  ")
    br.bold = True
    br.font.size = Pt(11)
    br.font.color.rgb = risk_color
    badge.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ── SECTION 1: Food Environment ───────────────────────────────
    supers      = int(g("supermarket", default=0))
    ff          = int(g("fast_food", default=0))
    conv        = int(g("convenience", default=0))
    ds          = int(g("dollar_store", default=0))
    swamp       = float(row.get("swamp_ratio", 0) or 0)
    ds_ratio    = float(row.get("dollar_store_ratio", 0) or 0)
    nearest     = float(row.get("nearest_supermarket_miles", np.nan))
    within_5    = int(row.get("supermarkets_within_5mi", 0))
    is_desert   = int(row.get("is_food_desert", 0))
    is_5mi      = int(row.get("is_desert_5mi", 0))
    is_mirage   = bool(row.get("is_food_mirage", False))
    is_r_mirage = bool(row.get("is_rural_mirage", False))

    food_rows = [
        ("Supermarkets in ZIP",
         str(supers) + (" ← no in-boundary supermarket" if supers == 0 else ""),
         "alert" if supers == 0 else "good"),
        ("Nearest supermarket",
         f"{nearest:.2f} miles" if not np.isnan(nearest) else "N/A",
         "alert" if nearest > 5 else "warn" if nearest > 2 else "good"),
        ("Supermarkets within 5 miles",
         str(within_5),
         "alert" if within_5 == 0 else "warn" if within_5 <= 2 else "good"),
        ("Fast food outlets",          str(ff),   "none"),
        ("Convenience stores",         str(conv),  "none"),
        ("Dollar stores",              str(ds),   "none"),
        ("Swamp ratio (unhealthy/supermarket)",
         f"{swamp:.1f}x" if supers > 0 else "N/A (no supermarket)",
         "alert" if swamp > 10 else "warn" if swamp > 3 else "good"),
        ("Dollar store density ratio",
         f"{ds_ratio:.2f}" if supers > 0 else "N/A",
         "warn" if ds_ratio >= 1.0 else "none"),
        ("Classified as food desert (ZIP boundary)",
         "YES" if is_desert else "No",
         "alert" if is_desert else "good"),
        ("Classified as food desert (5-mile buffer)",
         "YES" if is_5mi else "No",
         "alert" if is_5mi else "good"),
        ("Urban food mirage",
         "YES — nominal access, high structural need" if is_mirage else "No",
         "warn" if is_mirage else "none"),
        ("Rural thin-access zone",
         "YES" if is_r_mirage else "No",
         "warn" if is_r_mirage else "none"),
    ]
    _kv_section(doc, "1. Food Environment", food_rows)

    # ── SECTION 2: Demographics ───────────────────────────────────
    demo_rows = [
        ("Median household income",
         g("median_income", "${:,.0f}"),
         "none"),
        ("Population",
         g("population", "{:,.0f}"),
         "none"),
        ("Population density (per sq mi)",
         g("pop_density", "{:,.0f}"),
         "none"),
        ("Poverty rate",
         g("pct_poverty", "{:.1f}%"),
         "alert" if float(row.get("pct_poverty", 0) or 0) > 20
         else "warn" if float(row.get("pct_poverty", 0) or 0) > 10 else "none"),
        ("SNAP participation rate",
         g("pct_snap", "{:.1f}%"),
         "warn" if float(row.get("pct_snap", 0) or 0) > 15 else "none"),
        ("No-vehicle households",
         g("pct_no_vehicle", "{:.1f}%"),
         "warn" if float(row.get("pct_no_vehicle", 0) or 0) > 15 else "none"),
        ("Transit commuters",
         g("pct_transit", "{:.1f}%"),
         "none"),
        ("Elderly population (65+)",
         g("pct_elderly", "{:.1f}%"),
         "none"),
        ("College-educated adults",
         g("pct_college", "{:.1f}%"),
         "none"),
        (f"District property tax rate (2024)",
         g("district_tax_rate", "{:.3f}%"),
         "none"),
    ]
    _kv_section(doc, f"2. Demographics (ACS {acs_yr})", demo_rows)

    # ── SECTION 3: Economic Stress ────────────────────────────────
    stress_rows = [
        ("Economic stress score",
         g("economic_stress_score", "{:.3f}"),
         "none"),
        ("Need burden (poverty + SNAP)",
         g("need_burden", "{:.3f}"),
         "none"),
        ("Income stress index",
         g("income_stress", "{:.3f}"),
         "none"),
        ("Mobility stress index",
         g("mobility_stress", "{:.3f}"),
         "none"),
    ]
    _kv_section(doc, "3. Economic Stress Indicators", stress_rows)

    # ── SECTION 4: Health Outcomes ────────────────────────────────
    if health_row is not None:
        health_rows = [
            ("Diabetes prevalence",
             hg("DIABETES_CrudePrev", "{:.1f}%"),
             "warn" if float(health_row.get("DIABETES_CrudePrev", 0) or 0) > 12 else "none"),
            ("Obesity prevalence",
             hg("OBESITY_CrudePrev", "{:.1f}%"),
             "warn" if float(health_row.get("OBESITY_CrudePrev", 0) or 0) > 35 else "none"),
            ("Physical inactivity (no leisure activity)",
             hg("LPA_CrudePrev", "{:.1f}%"),
             "none"),
            ("Smoking prevalence",
             hg("CSMOKING_CrudePrev", "{:.1f}%"),
             "none"),
            ("Food insecurity prevalence",
             hg("FOODINSECU_CrudePrev", "{:.1f}%"),
             "warn" if float(str(health_row.get("FOODINSECU_CrudePrev", 0)).replace(",","") or 0) > 15 else "none"),
            ("SNAP participation (PLACES estimate)",
             hg("FOODSTAMP_CrudePrev", "{:.1f}%"),
             "none"),
            ("Poor mental health (≥14 days/month)",
             hg("MHLTH_CrudePrev", "{:.1f}%"),
             "none"),
            ("Depression prevalence",
             hg("DEPRESSION_CrudePrev", "{:.1f}%"),
             "none"),
        ]
        _kv_section(doc, "4. Health Outcomes (CDC PLACES 2024)", health_rows)
    else:
        _body(doc, "4. Health Outcomes — nj_health_places.csv not found. "
                   "Run the full pipeline to generate this file.", italic=True, color=C_GREY)

    # ── SECTION 5: Model Scores ───────────────────────────────────
    thresh = meta.get("best_threshold", 0.25)
    model_rows = [
        ("Desert probability score",
         f"{prob:.4f}  ({prob:.1%})",
         "alert" if prob >= 0.75 else "warn" if prob >= 0.40 else "good"),
        ("Predicted desert (default threshold 0.50)",
         "YES" if int(row.get("predicted_desert", 0)) else "No",
         "alert" if int(row.get("predicted_desert", 0)) else "none"),
        (f"Predicted desert (tuned threshold {thresh:.2f})",
         "YES" if int(row.get("predicted_desert_tuned", 0)) else "No",
         "alert" if int(row.get("predicted_desert_tuned", 0)) else "none"),
        ("Food access typology",
         str(typology),
         "alert" if typology == "Desert" else
         "warn"  if typology in ("Urban Mirage", "Rural Mirage", "Fragile Access") else
         "good"),
        ("Food mirage composite score",
         g("mirage_composite", "{:.3f}"),
         "none"),
    ]
    _kv_section(doc, "5. Model Outputs", model_rows)

    mirage_composite = float(row.get("mirage_composite", 0) or 0)
    swamp_ratio_val = float(row.get("swamp_ratio", 0) or 0)

    if typology == "Swamp" and mirage_composite > 0.70:
        _body(doc,
              f"Note: This ZIP is classified as a Swamp (swamp ratio {swamp_ratio_val:.1f}x) "
              f"but also scores in the {int(mirage_composite * 100):.0f}th percentile on the "
              f"mirage composite index, suggesting a dual food environment problem — "
              f"unhealthy food oversupply combined with structural access barriers for "
              f"elderly and mobility-limited residents.",
              color=C_AMBER
              )

    # ── SECTION 6: State Context ──────────────────────────────────
    _heading(doc, "6. State Context (Percentile Rankings)", size=11)

    context_items = [
        ("desert_probability",   "Desert probability",         False),
        ("nearest_supermarket_miles", "Distance to nearest supermarket", False),
        ("pct_poverty",          "Poverty rate",               False),
        ("swamp_ratio",          "Swamp ratio",                False),
        ("economic_stress_score","Economic stress score",      False),
        ("median_income",        "Median income",              True),
    ]
    lines = []
    for col, label, invert in context_items:
        val = row.get(col, np.nan)
        if pd.isna(val):
            continue
        val = float(val)
        series = pd.to_numeric(scores[col], errors="coerce").dropna()
        pct = (series < val).mean() * 100
        if invert:
            pct = 100 - pct
        arrow = "▲" if pct >= 75 else "▼" if pct <= 25 else "●"
        lines.append(f"  {arrow}  {label}: {pct:.0f}th percentile statewide")

    for line in lines:
        p = doc.add_paragraph()
        p.add_run(line).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # ── SECTION 7: Narrative Summary ─────────────────────────────
    doc.add_paragraph()
    _heading(doc, "7. Summary Assessment", size=11)

    # Build a dynamic narrative from the data
    narrative_parts = []

    if is_5mi:
        narrative_parts.append(
            f"ZIP code {zip_code} is a confirmed food desert — no supermarket "
            f"exists within five miles of the ZIP centroid. The nearest store is "
            f"{nearest:.1f} miles away."
        )
    elif is_desert:
        narrative_parts.append(
            f"ZIP code {zip_code} lacks a supermarket within its own boundaries, "
            f"but has {within_5} supermarket(s) within five miles. "
            f"The nearest is {nearest:.1f} miles from the ZIP centroid."
        )
    else:
        narrative_parts.append(
            f"ZIP code {zip_code} has {supers} supermarket(s) within its boundaries "
            f"and {within_5} within five miles. The nearest is {nearest:.1f} miles away."
        )

    if is_mirage:
        narrative_parts.append(
            "Despite nominal supermarket access, this ZIP is flagged as an urban food "
            "mirage — the need-to-capacity ratio places it in the top decile statewide, "
            "suggesting the single store may be insufficient for the community's needs."
        )


    poverty = float(row.get("pct_poverty", 0) or 0)
    no_car  = float(row.get("pct_no_vehicle", 0) or 0)
    if poverty > 15 or no_car > 15:
        narrative_parts.append(
            f"Economic vulnerability is elevated: poverty rate {poverty:.1f}% and "
            f"{no_car:.1f}% of households lack a vehicle, compounding access barriers."
        )

    if swamp > 5 and supers > 0:
        narrative_parts.append(
            f"The food environment is dominated by unhealthy options — the swamp ratio "
            f"of {swamp:.1f}x means fast food and convenience stores outnumber "
            f"supermarkets by more than five to one."
        )

    narrative_parts.append(
        f"The model assigns a desert probability of {prob:.1%}, classifying this "
        f"ZIP as {risk_label.lower()} for food access deprivation."
    )

    for part in narrative_parts:
        _body(doc, part)

    # In build_profile(), after the typology section:
    elderly_pct = float(row.get("pct_elderly", 0) or 0)
    transit_pct = float(row.get("pct_transit", 0) or 0)
    mobility = float(row.get("mobility_stress", 0) or 0)

    if elderly_pct > 22 and transit_pct < 3:
        _body(doc,
              f"⚠  Demographic Fragility Flag: {elderly_pct:.1f}% of residents are 65+ "
              f"(NJ avg ≈ 17%) with only {transit_pct:.1f}% transit commuters. "
              f"Formal supermarket proximity does not reflect practical access for a "
              f"population this age-skewed without reliable transit.",
              color=C_AMBER
              )

    # ── SECTION 8: Transit Access ─────────────────────────────────
    if gtfs:
        _heading(doc, "8. Transit Access Analysis", size=11)

        # Route 559
        r559 = gtfs.get("route_559", {})
        oc6 = gtfs.get("oc6", {})

        if r559:
            _kv_section(doc, "NJ Transit Route 559 (Fixed Route — In GTFS)", [
                ("Daily trips",
                 str(r559.get("trips_total", "—")),
                 "none"),
                ("Service span",
                 f"{r559.get('first_bus', '—')} – {r559.get('last_bus', '—')}",
                 "none"),
                ("Average headway",
                 f"{r559.get('avg_gap_min', '—')} minutes",
                 "none"),
                ("Nearest stop to Acme",
                 f"{r559.get('acme_stop', '—')} — "
                 f"{r559.get('acme_dist', '—')} mi "
                 f"({'✓ walkable' if r559.get('acme_walkable') else '✗ NOT walkable'})",
                 "good" if r559.get("acme_walkable") else "alert"),
                ("Nearest stop to Walmart",
                 f"{r559.get('walmart_stop', '—')} — "
                 f"{r559.get('walmart_dist', '—')} mi "
                 f"({'✓ walkable' if r559.get('walmart_walkable') else '✗ NOT walkable — highway crossing'})",
                 "good" if r559.get("walmart_walkable") else "alert"),
            ])

        if oc6:
            _kv_section(doc, "Ocean Ride OC 6 (County Route — NOT in GTFS)", [
                ("Service days",
                 oc6.get("service_days", "—"),
                 "warn"),
                ("No service",
                 "Friday, Saturday, Sunday",
                 "alert"),
                ("Bag limit",
                 f"{oc6.get('max_bags', '—')} bags maximum — hard constraint on weekly grocery volume",
                 "alert"),
                ("Senior fare",
                 oc6.get("fare_senior", "—"),
                 "good"),
                ("Wheelchair accessible",
                 "Yes",
                 "good"),
                ("Direct grocery stops",
                 ", ".join(oc6.get("stores_served", [])),
                 "good"),
                ("In NJ Transit GTFS",
                 "NO — invisible to USDA, proximity models, and this ML pipeline",
                 "alert"),
                ("AM shopping window — Walmart",
                 f"Arrive {oc6.get('am_walmart_arrive', '—')}, "
                 f"board by {oc6.get('am_walmart_depart', '—')} "
                 f"({oc6.get('am_walmart_window', '—')} min)",
                 "none"),
                ("PM shopping window — Walmart",
                 f"Arrive {oc6.get('pm_walmart_arrive', '—')}, "
                 f"board by {oc6.get('pm_walmart_depart', '—')} "
                 f"({oc6.get('pm_walmart_window', '—')} min — TIGHT)",
                 "warn"),
            ])

            p = doc.add_paragraph()
            r = p.add_run(
                "⚠  Critical Gap: OC 6 is the only route providing door-to-door grocery "
                "access for seniors and car-free residents in this ZIP. Its absence from "
                "GTFS means it does not appear in the USDA Food Access Research Atlas, "
                "NJ Transit trip planners, or this ML model. The 3-day service constraint "
                "and 2-bag limit structurally cap weekly food provisioning capacity."
            )
            r.font.size = Pt(10)
            r.font.color.rgb = C_AMBER
            r.bold = True
            p.paragraph_format.space_after = Pt(8)
    else:
        # No GTFS data — add a placeholder note
        _heading(doc, "8. Transit Access", size=11)
        _body(doc,
              f"No GTFS analysis file found for ZIP {zip_code}. "
              f"Run GTFSNJ.py for this ZIP to include transit access data.",
              italic=True, color=C_GREY
              )

    # ── Footer ────────────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f"ZIP {zip_code} Profile  |  {today}  |  "
        f"Model: Random Forest (AUC {meta.get('test_auc', '—')})  |  "
        f"ACS {acs_yr} · OSM · CDC PLACES 2024 · USDA FARA {meta.get('usda_year', 2019)}"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = C_GREY
    fr.italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[✓] Profile saved → {out_path}")


def main():
    zip_code = input("Enter NJ ZIP code: ").strip().zfill(5)

    if not zip_code.isdigit() or len(zip_code) != 5:
        print(f"[ERROR] '{zip_code}' is not a valid ZIP code.")
        return

    if not zip_code.startswith(("07", "08")):
        print(f"[WARNING] '{zip_code}' may not be a New Jersey ZIP code.")
        proceed = input("Continue anyway? (y/n): ").strip().lower()
        if proceed != "y":
            return

    out_path = HERE / "outputs" / f"{zip_code}_food_access_profile.docx"
    print(f"Generating profile for {zip_code} → {out_path}")

    try:
        build_profile(zip_code, out_path)
    except ValueError as e:
        print(f"[ERROR] {e}")
if __name__ == "__main__":
    main()
